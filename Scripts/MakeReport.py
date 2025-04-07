import os
import yaml
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
import sys
import numpy as np
import cv2
import shutil
try:
    from Scripts.ReportUtils import *
except:
    from ReportUtils import *

type_locations = ["Hầu họng", "Thực quản", "Tâm vị", "Thân vị", "Phình vị",
                  "Hang vị", "Bờ cong lớn", "Bờ cong nhỏ", "Hành tá tràng", "Tá tràng"]
type_diseases = ["Viêm thực quản trào ngược", "Ung thư thực quản",
                 "Viêm dạ dày", "Ung thư dạ dày", "Loét hành tá tràng", "Ton thuong"]
label_det_dict = {
    "Viêm thực quản": "Viêm thực quản trào ngược",
    "Viêm thực quản trào ngược": "Viêm thực quản trào ngược",
    "Viêm dạ dày HP": "Viêm dạ dày",
    "Viêm dạ dày": "Viêm dạ dày",
    "Ung thư dạ dày": "Ung thư dạ dày",
    "Ung thư thực quản": "Ung thư thực quản",
    "Loét hành tá tràng": "Loét hành tá tràng",
    "Ton thuong": "Ton thuong"
}
correct = "☒"
incorrect = "☐"


import subprocess
import os
import platform

def open_docx(file_path):
    # Check if the file exists
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The file '{file_path}' does not exist.")
    
    # Get the operating system
    os_name = platform.system()

    if os_name == 'Windows':
        # Use start command for Windows
        subprocess.run(['start', '', file_path], shell=True)
    elif os_name == 'Linux':
        # Use xdg-open command for Linux
        subprocess.run(['xdg-open', file_path])
    elif os_name == 'Darwin':
        # Use open command for macOS
        subprocess.run(['open', file_path])
    else:
        raise OSError(f"Unsupported operating system: {os_name}")

def open_folder(folder_path):
    # Check if the folder exists
    if not os.path.isdir(folder_path):
        raise FileNotFoundError(f"The folder '{folder_path}' does not exist.")
    
    # Get the operating system
    os_name = platform.system()

    if os_name == 'Windows':
        # Use explorer command for Windows
        subprocess.run(['explorer', folder_path])
    elif os_name == 'Linux':
        # Use xdg-open command for Linux
        subprocess.run(['xdg-open', folder_path])
    elif os_name == 'Darwin':
        # Use open command for macOS
        subprocess.run(['open', folder_path])
    else:
        raise OSError(f"Unsupported operating system: {os_name}")

def make_report(folder, config, hospitalCode, openWord=True, openExplorer=True):
    minConfidence = config['minConfidence']
    outDir = config['outDir']
    docTemplate = config['docTemplate']
    similarityThreshold = get_value(config,'similarityThreshold',0.3)
    numImages = get_value(config,'numImages',8)
    if not os.path.exists(outDir):
        os.makedirs(outDir)
    

    doc_path = os.path.join('Resources', 'Report', docTemplate)
    doc = Document(doc_path)

    outDirFolder = os.path.join(outDir, hospitalCode, folder.name)
    if not os.path.exists(outDirFolder):
        os.makedirs(outDirFolder)

    
    outDirFolderGiaiPhau = os.path.join(outDirFolder,'1. Vi tri giai phau')
    if not os.path.exists(outDirFolderGiaiPhau):
        os.makedirs(outDirFolderGiaiPhau)

    
    outDirFolderTonThuong = os.path.join(outDirFolder,'2. Ton thuong')
    if not os.path.exists(outDirFolderTonThuong):
        os.makedirs(outDirFolderTonThuong)

    out_path_doc = os.path.join(outDirFolder, "report.docx")
    folder.parse_report()
    for image in folder.images:
        image.parse_yaml()
    patient = folder.patient
    patient_details = {
            "Mã bệnh nhân:": patient.id,
            "Tên bệnh nhân:": patient.name,
            "Bác sĩ nội soi:": patient.doctor,
            "Thời gian bắt đầu": patient.start_time,
            "Thời gian kết thúc": patient.end_time,
            "Ghi chú:": patient.info,
            "Giới tính:": patient.gender,
            "Năm sinh:": patient.birth_year,
        }
    
    # Iterate through the paragraphs to fill in the patient details
    for paragraph in doc.paragraphs:
            for key, value in patient_details.items():
                if key in paragraph.text:
                    paragraph.text = key + " " + str(value)
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = docx.shared.Pt(14)
    
    type_locations = ["Hầu họng", "Thực quản", "Tâm vị", "Thân vị", "Phình vị",
                  "Hang vị", "Bờ cong lớn", "Bờ cong nhỏ", "Hành tá tràng", "Tá tràng"]
    type_diseases = ["Viêm thực quản trào ngược", "Ung thư thực quản",
                 "Viêm dạ dày", "Ung thư dạ dày", "Loét hành tá tràng"]
    title1 = 'Tỷ lệ các vị trí giải phẫu xuất hiện đầy đủ: '
    title2 = 'Tổn thương phát hiện được:'
    table_data = {
        "Vị trí": {
        },
        "Loại tổn thương": {
        }
    }
    disease_counts = {}
    for image_obj in folder.images:
        obj = table_data['Vị trí']
        if image_obj.regionProbability >= minConfidence:
            region = image_obj.region
            if region is not None:
                if region not in disease_counts: disease_counts[region] = []
                disease_counts[region].append([image_obj, None])
        for box in image_obj.boxes:
            if box.conf >= minConfidence:
                label = box.label
                if label not in disease_counts: disease_counts[label] = []
                disease_counts[label].append([image_obj,box])
    for k in disease_counts:
        print (k, len(disease_counts[k]))
    for labels, name, title in zip([type_locations, type_diseases], ['Vị trí', 'Loại tổn thương'], [title1, title2]):
        total = 0
        countExists = 0
        for label in labels:
            count = 0
            save_d = remove_accents(label).lower().replace(' ', '_')
            if label in disease_counts: count = len(disease_counts[label])
            total += 1
            icon = correct if count>0 else incorrect
            countExists += 1 if count>0 else 0
            table_data[name][label] = {"detected": icon, "link": save_d, "count": str(count)}
        for paragraph in doc.paragraphs:
            if title in paragraph.text:
                paragraph.text = title + f" {countExists}/{total}"
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = docx.shared.Pt(14)

    for table in doc.tables:
        header = table.cell(0, 0).text.strip()
        if header in table_data:
            update_table(table, table_data[header])
        else:
            pass
    
    # print ('table_data',table_data)

    #save and insert images

    for label in type_locations:
        label_formatted = remove_accents(label).lower().replace(' ', '_')
        save_d = os.path.join(outDirFolderGiaiPhau, label_formatted)
        if label in disease_counts:
            for image_obj, _ in disease_counts[label]:
                if not os.path.exists(save_d): os.makedirs(save_d)
                out_path = os.path.join(save_d, image_obj.base_name)
                shutil.copy(image_obj.file_path, out_path)

    image_dict = []
    for label in type_diseases:
        label_formatted = remove_accents(label).lower().replace(' ', '_')
        save_d = os.path.join(outDirFolderTonThuong, label_formatted)
        if label in disease_counts:
            for image_obj, box in disease_counts[label]:
                image = cv2.imread(image_obj.file_path)
                histogram = compute_hsv_histogram(image)
                cv2.rectangle(image, (box.x,box.y,box.width,box.height),
                                (100, 255, 100), 10)
                score_str = str(box.conf*100)[:5]+'%'
                box_id = image_obj.boxes.index(box)
                cv2.putText(image, score_str, (int(
                    box.x+30), int(box.y+60)), cv2.FONT_HERSHEY_DUPLEX, 1.5, (0, 255, 0), 7)

                if not os.path.exists(save_d): os.makedirs(save_d)
                out_path = os.path.join(save_d, image_obj.base_name[:-4]+'_'+str(box_id)+'.jpg')
                cv2.imwrite(out_path, image)
                image_dict.append([out_path, label, box, image.shape, histogram, image_obj.selected])
                # print ("histogram",histogram, image_obj.selected)
                # shutil.copy(image_obj.file_path, out_path)


    for idx, label in enumerate(type_diseases):
            arr = []
            for path, label2, box, shape, histogram, selected in image_dict:
                h, w = shape[:2]
                if label == label2:
                    arr.append([path, box.conf, w, h, histogram, selected])
            arr = select_top_k_images(arr, numImages, similarityThreshold)
            # arr.sort(key=lambda t: t[1], reverse=True)
            arr = arr[:numImages]
            if len(arr) > 0:
                # images = [a[0] for a in arr]
                title = f"3.{idx+1} {label}"
                insert_images_in_grid(doc, title, arr)
    try:
        doc.save(out_path_doc)
        if openWord:
            open_docx(out_path_doc)
    except: pass
    if openExplorer:
        open_folder(outDirFolder)


s1 = u'ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝàáâãèéêìíòóôõùúýĂăĐđĨĩŨũƠơƯưẠạẢảẤấẦầẨẩẪẫẬậẮắẰằẲẳẴẵẶặẸẹẺẻẼẽẾếỀềỂểỄễỆệỈỉỊịỌọỎỏỐốỒồỔổỖỗỘộỚớỜờỞởỠỡỢợỤụỦủỨứỪừỬửỮữỰựỲỳỴỵỶỷỸỹ'
s0 = u'AAAAEEEIIOOOOUUYaaaaeeeiioooouuyAaDdIiUuOoUuAaAaAaAaAaAaAaAaAaAaAaAaEeEeEeEeEeEeEeEeIiIiOoOoOoOoOoOoOoOoOoOoOoOoUuUuUuUuUuUuUuYyYyYyYy'

def remove_accents(input_str):
    s = ''
    for c in input_str:
        if c in s1:
            s += s0[s1.index(c)]
        else:
            s += c
    return s


def add_hyperlink_to_cell(cell, url, text):
    # Clear the current content in the cell
    cell.text = ""

    # Create the w:hyperlink tag and add needed attributes
    part = cell.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if you want to (optional)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')  # Hex color for blue
    rPr.append(c)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Create a w:t element and add the text to it
    text_elem = OxmlElement('w:t')
    text_elem.text = text

    # Add the w:rPr and w:t elements to the w:r element
    new_run.append(rPr)
    new_run.append(text_elem)

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the hyperlink to the cell
    cell._element.append(hyperlink)


def update_table(table, data):
    for row in table.rows:
        cells = row.cells
        key = cells[0].text.strip()
        if key in data:
            cells[1].text = data[key]["detected"]
            # cells[2].text = data[key]["link"]
            cells[2].text = data[key]["count"]
            for cell in [cells[1], cells[2]]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = docx.shared.Pt(14)

            # add_hyperlink_to_cell(cells[2], 'https://google.com',data[key]["link"])


def insert_images_in_grid(doc, title, images, grid_size=(4, 2)):

    from docx.shared import Inches
    from docx.shared import Inches, Pt
    # Determine grid dimensions
    rows, cols = grid_size

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(14)  # Set the font size for the title
    title_run.font.name = 'Times New Roman'

    table = doc.add_table(rows=rows, cols=cols)

    # Insert images into the table
    img_index = 0
    for row in range(rows):
        for col in range(cols):
            if img_index < len(images):
                # Get the cell where the image should be placed
                cell = table.cell(row, col)

                # Add a new paragraph to the cell and insert the image
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                image_path, best_score, w, h,_,_ = images[img_index]
                dst_height = 3*h/w
                run.add_picture(image_path, width=Inches(3), height=Inches(
                    dst_height))  # Adjust width and height as needed

                img_index += 1

    # title_paragraph = doc.add_paragraph()
    # title_run = title_paragraph.add_run('')
    # title_run.font.size = Pt(11)  #

# example: python .\MakeReport.py C:\Users\HVN\Desktop\endo\app_noi_soi\utils\logs\106_22-06-2024-10-01-57 out
if __name__ == "__main__":
    indir = sys.argv[1]
    outdir = sys.argv[2]
    min_confidence = 0.8
    if len(sys.argv) > 3:
        min_confidence = float(sys.argv[3])
    script_dir = os.path.dirname(os.path.abspath(__file__))
    config = {
        "minConfidence":min_confidence,
        "outDir":outdir,
        "docTemplate": os.path.join(script_dir,"maubaocao4.docx")
    }
    hospitalCode = ''
    session_path = os.path.join(indir, "session.yaml")
    if os.path.exists(session_path):
        folderName = os.path.basename(indir)

        folder = Folder(indir,folderName,None,None,None)
        folder.parse_report()

        files = os.listdir(indir)
        for f in files:
            f_name = f.split('.')[0]
            path = os.path.join(indir, f)
            yaml_path = os.path.join(indir, f_name+'.yaml')
            if path.endswith('jpg') and os.path.exists(yaml_path):
                folder.images.append(ReportImage(path))
            
        # print (folderName)
        make_report(folder, config, hospitalCode)
    else:
        dirs = os.listdir(indir)
        for folderName in dirs:
            indir2 = os.path.join(indir, folderName)
            print ("indir2",indir2)
            session_path = os.path.join(indir2, "session.yaml")
            if os.path.exists(session_path):
                folder = Folder(indir2,folderName,None,None,None)
                folder.parse_report()
                files = os.listdir(indir2)
                for f in files:
                    f_name = f.split('.')[0]
                    path = os.path.join(indir2, f)
                    yaml_path = os.path.join(indir2, f_name+'.yaml')
                    if path.endswith('jpg') and os.path.exists(yaml_path):
                        folder.images.append(ReportImage(path))
                    
                print (folderName)
                make_report(folder, config, hospitalCode, False, False)
        open_folder(outdir)
